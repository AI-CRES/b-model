import streamlit as st
import openai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import datetime
import re
from bs4 import BeautifulSoup


hide_streamlit_style = """
    <style>
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    </style>
    """
st.markdown(hide_streamlit_style, unsafe_allow_html=True)

# Configuration de l'API OpenAI

api_key = st.secrets["API_KEY"]
openai.api_key = api_key

# ----------------------------------------------------------------------------
# 1) Fonction pour appeler ChatGPT et générer le texte du Business Model Canvas
# ----------------------------------------------------------------------------
def get_metaprompt(type_entreprise):
    """
    Retourne un metaprompt spécifique basé sur le type d'entreprise.
    """
    metaprompts = {
        "PME": """**Méta-Prompt pour l’Élaboration d’un Business Model pour PME Traditionnelle (Intégrant des Innovations Low-Tech et Adaptées aux Contextes Africains ou Émergents)**

        **Votre Rôle :**  
        Vous êtes un expert en stratégie d’entreprise, marketing, UX, innovation frugale (low-tech et éventuellement high-tech), et élaboration de Business Models. Vous devez générer un Business Model complet, clair, chiffré, cohérent et innovant, adapté à une PME qui opère dans un environnement local (par exemple en Afrique ou dans d’autres pays émergents) où les réalités technologiques, économiques, culturelles et réglementaires diffèrent des contextes occidentaux fortement numérisés.  
        L’innovation ne sera pas seulement technologique de pointe (high-tech), mais aussi low-tech (solutions simples, robustes, faciles d’entretien, peu consommatrices de ressources), et tenant compte des infrastructures limitées, des préférences culturelles, de la disponibilité intermittente de l’électricité, du coût de la connectivité, de l’importance du lien social, etc.

        Votre tâche s’organise en trois phases :  
        1. Configuration Initiale (Collecte et Structuration des Données)  
        2. Étapes Intermédiaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)  
        3. Production Finale (Business Model Canvas)

        Pour chaque phase, suivez les instructions et veillez à :  
        - Prendre en compte la persona (données démographiques, comportementales, capacités d’adoption de l’innovation)  
        - Analyser le marché (taille, segments, offres existantes formelles et informelles, niveau de satisfaction, tendances locales, disponibilité de ressources, logistique)  
        - Intégrer les facteurs limitants (technologiques, économiques, culturels, psychologiques, physiologiques, réglementaires, infrastructures limitées)  
        - Évaluer la concurrence (locale, informelle, substituts traditionnels), comprendre les niveaux de satisfaction et de confiance  
        - Comprendre le parcours client (avant, pendant, après), intégrer la carte d’empathie, identifier les gains et souffrances spécifiques au contexte (par exemple : importance du bouche-à-oreille, confiance interpersonnelle, exigence de robustesse, maintenance locale)  
        - Vérifier systématiquement la cohérence, proposer des optimisations et ajustements  
        - Avant d’introduire une innovation (low-tech ou high-tech), s’assurer que la persona est prête à l’adopter, en tenant compte de l’accessibilité, du coût, de la simplicité et de la réputation  
        - Produire un Business Model Canvas complet (9 blocs), avec des méta-prompts spécifiques pour chacun des blocs, adaptés au contexte local

        ---

        ### Phase 1 : Configuration Initiale (Entrée de Données)

        1. **Recueille et structure les informations suivantes :**  
        - **Persona** :  
            - Données démographiques : âge, sexe, localisation (zones urbaines, péri-urbaines, rurales), niveau d’éducation (alphabétisation, langues parlées), profession (artisans, commerçants, agriculteurs, employés, indépendants), revenu moyen.  
            - Paramètres comportementaux : sensibilité au prix (budgets limités, nécessité de micro-paiements), fréquence et volume d’achat (achats ponctuels, saisonniers, hebdomadaires), perception de la qualité (fiabilité, durabilité), utilisation technologique (téléphones basiques, smartphones d’entrée de gamme, accès limité à Internet), accessibilité (distance aux points de vente, transport limité), temps disponible (horaires de travail, saison des récoltes), besoins spécifiques (ex : accès à l’eau, énergie, outils agricoles, services financiers de base, éducation des enfants, soins de santé).  
            - Capacité d’adoption de l’innovation : Familiarité avec certaines technologies (mobile money, radios communautaires, solutions solaires), ouverture au changement dépendant de la preuve sociale, de la confiance dans la communauté, de la simplicité et robustesse du produit/service. Barrières psychologiques/culturelles (méfiance envers les nouvelles solutions étrangères, préférence pour le contact humain, importance de la recommandation de la famille ou du chef de village).  
        
        - **Analyse du Marché** :  
            - Taille du marché local : estimer la population concernée, le pouvoir d’achat moyen, les infrastructures disponibles.  
            - Segments : populations urbaines vs rurales, artisans, commerçants, coopératives, PME locales, secteur informel.  
            - Offres concurrentes existantes : solutions traditionnelles (artisanales, informelles), importations bas de gamme, programmes d’ONG, concurrents locaux ou étrangers, modèles low-cost.  
            - Niveau de satisfaction actuel : Les clients sont-ils satisfaits des solutions actuelles ? Y a-t-il un manque de fiabilité, de formation, de SAV ?  
            - Tendances : adoption progressive du mobile money, sensibilisation croissante à l’énergie solaire, émergence de petites coopératives, engouement pour des solutions durables et réparables.  
            - Innovations et comportements émergents : réemploi, économie circulaire, mise en commun de ressources, augmentation des transferts d’argent via mobile.  
        
        - **Facteurs Limitants** :  
            - Contraintes technologiques : faible accès à l’électricité stable, couverture internet inégale, outils technologiques rudimentaires, importance de solutions low-tech (pompes manuelles, panneaux solaires simples, systèmes de filtration d’eau basiques).  
            - Contraintes économiques : revenus limités, volatilité des prix, accès restreint au crédit, nécessité d’étaler les paiements (micro-paiements, crédit rotatif, tontines).  
            - Contraintes culturelles : langues locales, importance de la confiance interpersonnelle, réticence à adopter des produits inconnus sans démonstration ou validation par la communauté.  
            - Contraintes psychologiques et physiologiques : besoin de solutions simples d’utilisation, ergonomiques, adaptées aux conditions climatiques (chaleur, poussière), faible taux d’alphabétisation nécessitant des modes d’emploi visuels.  
            - Contraintes réglementaires : normes locales, barrières douanières, absence de normes formelles dans certains secteurs, difficulté à obtenir des certifications officielles.  

        **Après avoir recueilli ces données, effectue une première analyse critique** :  
        - Vérifie la cohérence des informations.  
        - Identifie les lacunes (par exemple, manque d’informations sur le pouvoir d’achat réel, sur le réseau de distribution informel, sur le rôle des leaders d’opinion locaux).  
        - Propose des compléments ou ajustements pour optimiser la qualité des données (ajouter des données sur la saisonnalité du marché, l’influence des ONG, l’impact des conditions climatiques, la présence ou non de microfinance).

        ---

        ### Phase 2 : Étapes Intermédiaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)

        2. **Analyse du Parcours Client & Carte d’Empathie** :  
        - Décris le parcours client (avant, pendant, après) en tenant compte des conditions locales :  
            - Avant : Le client prend conscience de son besoin par le bouche-à-oreille, via un ami, un voisin, un leader communautaire, ou en écoutant la radio locale. Il compare avec les solutions déjà connues (artisan local, réparations informelles, solutions importées). Il évalue la confiance, le prix, la disponibilité.  
            - Pendant : Achat sur un marché local, essai d’une démonstration concrète (démonstration en conditions réelles, sur une place de village), informations données par un vendeur itinérant ou un agent de confiance. Utilisation d’un mode de paiement adapté (cash, mobile money).  
            - Après : Suivi du produit, entretien, besoin de pièces détachées, possibilité de contact direct avec l’entreprise (ligne téléphonique, point de service local), échange d’expériences avec d’autres utilisateurs, éventuel SAV simplifié (réparations locales, pièces détachées bon marché).  
        
        - Identifie les points de contact (marchés, boutiques, intermédiaires locaux, radios communautaires, SMS informatifs), obstacles (faible connectivité, manque d’informations détaillées, barrières linguistiques), moments de vérité (premier essai du produit, première panne et réactivité du SAV), frustrations (produit pas adapté, manuel incompréhensible, manque de fiabilité).  
        
        - Intègre les contraintes physiologiques, psychologiques, économiques, culturelles, technologiques, réglementaires : par exemple, l’importance de la simplicité et de la robustesse pour réduire la crainte d’une technologie trop complexe, la nécessité de support en langue locale, la possibilité de s’adapter aux normes informelles.  
        
        - Crée une carte d’empathie :  
            - Pensées : « Est-ce que cette solution est fiable, reconnue par ma communauté ? Est-ce que je vais perdre mon argent si ça ne marche pas ? »  
            - Sentiments : Méfiance, curiosité, besoin de réassurance, fierté s’il s’agit d’une innovation locale valorisée.  
            - Actions : Demande de conseils à d’autres, observation d’exemples concrets, volonté d’essayer avant d’acheter.

        3. **Gains et Souffrances** :  
        - Liste les gains : par exemple, accès facilité à un service vital (eau, énergie, outil de gestion commerciale simple), réduction du temps et de l’effort, robustesse (moins de pannes), accès à un SAV local, meilleure rentabilité ou productivité.  
        - Liste les souffrances : manque de solutions adaptées, problèmes de maintenance, coûts initiaux trop élevés sans option de paiement flexible, manque de formation pour utiliser correctement le produit.

        4. **Élaboration de la Carte de Valeur** :  
        - Définis la mission de consommation principale : répondre à un besoin fondamental (ex : un outil agricole robuste, une solution d’éclairage solaire fiable, un service financier simple via mobile, un appareil domestique low-tech adapté aux pannes d’électricité).  
        - Identifie les gains déjà fournis par les offres actuelles (ex : disponibilité locale, prix bas) et les souffrances non adressées (faible qualité, pas de SAV, pas d’adaptation aux conditions réelles).  
        - Esquisse une proposition de valeur préliminaire adaptée à la capacité d’adoption de l’innovation par la persona :  
            - Une solution simple, robuste, facilement compréhensible, qui peut être testée avant achat.  
            - Un modèle de distribution local (agents sur le terrain), un SAV accessible, un support en langue locale, des options de paiement flexible (mobile money, tontines, microcrédit).  
            - Intégration progressive d’innovations low-tech (p. ex. appareils mécaniques robustes, panneaux solaires portables) ou high-tech simple (SMS, USSD, application mobile légère) si l’utilisateur est prêt.

        5. **Détermination du Segment de Clients** :  
        - Choisis le type de relation (B2C direct, B2B via des coopératives, B2B2C via des distributeurs locaux).  
        - Priorise les segments qui correspondent le mieux :  
            - Par exemple, petits commerçants urbains ayant un pouvoir d’achat limité mais stables, agriculteurs nécessitant un outil fiable en milieu rural, coopératives d’artisans prêts à adopter une solution pour améliorer leur productivité.  
        - Tient compte de leur sensibilité au prix, de leur ouverture à l’innovation, de leur capacité à comprendre et utiliser la solution, de la nécessité de formation.

        6. **Analyse des Problèmes et Solutions (Canvas de Problème)** :  
        - Identifie les problèmes majeurs : par exemple, la difficulté à accéder à un produit fiable, le manque d’informations, la complexité du produit, le coût trop élevé d’une solution importée haut de gamme.  
        - Associe chaque problème à une solution :  
            - Problème : manque de SAV → Solution : réseau de réparateurs locaux formés.  
            - Problème : prix élevé d’entrée → Solution : offres en micro-paiements, location-vente, partenariats avec microfinance.  
            - Problème : manque de confiance → Solution : démonstrations, témoignages de pairs, communication via radios locales et leaders d’opinion.  
        - Justifie en quoi les solutions sont meilleures que l’existant : plus adaptées, plus abordables, plus simples, prenant en compte la réalité du terrain (faible infrastructure, besoin de résilience, faible taux d’alphabétisation).

        **Après ces étapes, fais une analyse intermédiaire** :  
        - Vérifie la cohérence du contexte, du parcours client, des solutions proposées.  
        - Assure-toi que les innovations (low-tech, partenariats locaux, solutions de paiement flexible) sont compréhensibles et adoptables par la persona.  
        - Propose des ajustements stratégiques : simplification du produit, ajustement du prix, ajout d’un canal de distribution plus local, formation des utilisateurs, partenariats avec des ONG ou des radios locales.

        ---

        ### Phase 3 : Production Finale du Business Model (Business Model Canvas)

        Sur la base des analyses précédentes, génère un Business Model Canvas complet. Utilise les méta-prompts suivants pour chaque bloc, en tenant compte du contexte local, des solutions low-tech et des infrastructures limitées :

        1. **Segments de Clients**  
        Méta-Prompt :  
        « Définis précisément les segments de clients ciblés, en tenant compte :  
        - De leurs caractéristiques sociodémographiques (âge, sexe, localisation, niveau d’éducation, profession, revenu, langue).  
        - De leurs comportements d’achat (fréquence, volume, sensibilité au prix, recours au crédit informel, canaux de confiance : marchés locaux, revendeurs informels, chefs de village, radios).  
        - De leur maturité technologique (téléphones basiques, usage de SMS/USSD, familiarité avec le mobile money, radio, bouche-à-oreille, rencontres physiques).  
        - De leur capacité d’adoption de l’innovation (ouverture au changement si démonstration concrète, barrières culturelles, besoin de preuves, préférences pour du low-tech robuste plutôt que du high-tech fragile).  
        - De leurs contraintes (faible pouvoir d’achat, saisons de récolte, temps de disponibilité, accès difficile à l’électricité ou à internet).  
        Intègre également des scénarios évolutifs (si l’économie se dégrade, réduction de l’achat ou passage à des solutions plus frugales ; si la technologie progresse, adoption graduelle de services numériques simples).  
        Justifie pourquoi ces segments sont retenus : potentiel de rentabilité, facilité d’accès via des canaux locaux, réceptivité à la proposition de valeur (améliorer leur vie de façon concrète, fiable, abordable). »

        2. **Proposition de Valeur**  
        Méta-Prompt :  
        « Détaille la proposition de valeur en explicitant :  
        - Les besoins fondamentaux (eau, énergie, information, outils productifs, services financiers simples).  
        - Les souffrances clientes (manque de fiabilité, difficulté d’entretien, complexité des produits, méfiance) et comment elles sont résolues (simplicité, robustesse, support local, preuves sociales).  
        - Les gains fournis (amélioration de la productivité, économies de temps, durabilité, réduction de la dépendance à des systèmes complexes, meilleure gestion financière) et inclure les bénéfices émotionnels (confiance, fierté, reconnaissance sociale).  
        - La différenciation par rapport aux offres concurrentes : intégration dans le tissu local, formation d’agents locaux, facilité d’entretien, pricing adapté, low-tech combiné avec technologie simple (mobile money), SAV local.  
        - L’introduction progressive de l’innovation : démonstrations pratiques, formation sur le terrain, tutoriels en langue locale, partenariat avec leaders communautaires.  
        - Variantes selon les segments : option premium (un meilleur SAV, une maintenance plus poussée) pour les clients plus solvables, version ultra-simplifiée pour les segments plus conservateurs ou à très faible pouvoir d’achat. »

        3. **Canaux de Distribution**  
        Méta-Prompt :  
        « Définis les canaux par lesquels les clients seront informés, convaincus, achèteront et utiliseront le produit/service, en tenant compte des réalités locales :  
        - Canaux hors ligne : marchés locaux, boutiques physiques, vente itinérante, radios communautaires, affichages, démonstrations sur place, coopératives agricoles, leaders religieux ou communautaires.  
        - Canaux digitaux légers : SMS, USSD, appels téléphoniques, WhatsApp, Facebook local, mobile money.  
        - Nécessité d’omnicanalité adaptée au contexte : cohérence entre communication radio, démonstration physique, et suivi par téléphone.  
        - Simplicité d’accès et besoin d’accompagnement pédagogique (formation dans les marchés, brochures visuelles, tutoriels audio).  
        - Adaptabilité des canaux si le marché évolue (ex: introduction progressive d’une application mobile si la connectivité s’améliore).  
        Justifie chaque canal (coût, accessibilité, confiance) et comment il réduit les obstacles à l’adoption, améliore la satisfaction, et s’intègre dans le parcours client local. »

        4. **Relations Clients**  
        Méta-Prompt :  
        « Décris la nature et la qualité des relations établies avec les clients :  
        - Personnalisation via un réseau d’agents locaux qui connaissent la langue, la culture, et les besoins.  
        - Communauté : création de groupes d’utilisateurs, d’associations locales, de rencontres de démonstration, événements communautaires où les clients échangent leurs expériences.  
        - Automatisation : mise en place d’un service SMS de rappel, d’une hotline téléphonique simple, d’un chatbot vocal si la technologie le permet (ou service d’appels humains en langue locale).  
        - Fidélisation : réductions pour clients fidèles, options de maintenance préventive, accès à des mises à jour techniques simples, partenariats avec des ONG pour aider à la formation continue.  
        - Gestion des plaintes et retours : politique claire de SAV, réparation locale, garantie adaptée, délais de réponse rapides.  
        Intègre la dimension culturelle (contact humain valorisé), psychologique (confiance, besoin de réassurance), réglementaire (respect des règles locales, si existantes). Explique comment ces relations évoluent au fil du temps et renforcent la CLV dans un contexte de marché volatile. »

        5. **Sources de Revenus**  
        Méta-Prompt :  
        « Détaille les mécanismes de génération de revenus :  
        - Modèles de tarification : vente directe à prix abordable, options de micro-paiements échelonnés, crédit via partenaire de microfinance, location-vente, abonnement léger (maintenance), freemium (démonstration gratuite, paiement pour les pièces détachées).  
        - Justification des prix : aligner le prix sur le pouvoir d’achat, offrir un excellent rapport qualité/durabilité/prix, tenir compte des référentiels locaux (si les concurrents informels sont très bon marché, justifier la valeur par la fiabilité).  
        - Réductions des freins économiques : essai avant achat, garantie satisfait ou remboursé, partenariats avec ONG ou institutions locales.  
        - Diversification des revenus : ventes croisées (pièces détachées, formation), partenariats B2B (ventes en gros à des coopératives), publicité locale, sponsorisation par des institutions de développement.  
        - Adaptation aux changements : si le marché se contracte, proposer des modèles encore plus frugaux, si la réglementation change, s’adapter par des produits conformes.  
        Explique comment cette structure de revenus soutient la viabilité à long terme et reste cohérente avec la proposition de valeur et la sensibilité au prix de la persona. »

        6. **Ressources Clés**  
        Méta-Prompt :  
        « Identifie toutes les ressources indispensables :  
        - Ressources Humaines : agents locaux (formés aux langues et contextes locaux), réparateurs, formateurs, personnels de SAV.  
        - Ressources Technologiques : outils de communication simples (téléphones basiques, logiciels légers), systèmes de paiement mobile, éventuellement une plateforme centralisée mais légère.  
        - Ressources Intellectuelles : savoir-faire sur l’adaptation du produit au contexte local, guides visuels, partenariats de R&D avec des instituts techniques locaux.  
        - Ressources Matérielles : pièces détachées robustes, matériaux durables, équipements simples qui ne nécessitent pas une infrastructure complexe.  
        - Ressources Financières : capital initial, fonds de roulement, accès à la microfinance ou à des investisseurs sociaux, trésorerie pour faire face aux saisons difficiles.  
        - Ressources Relationnelles : liens solides avec les communautés, chefs traditionnels, radios locales, ONG, institutions de développement.  
        Pour chaque ressource, justifie pourquoi elle est critique (ex. sans agents locaux, pas de confiance ; sans matériaux robustes, produit inutilisable), et comment ces ressources assurent un avantage concurrentiel durable. »

        7. **Activités Clés**  
        Méta-Prompt :  
        « Décris les activités indispensables :  
        - Développement & Innovation : adapter le produit aux conditions locales (climat, langue), améliorer la durabilité, simplifier l’usage.  
        - Production & Livraison : fabrication locale ou semi-locale, contrôle de la qualité, approvisionnement en pièces robustes, logistique simple (transport par camions, motos, ânes si nécessaire).  
        - Marketing & Ventes : communication via radios communautaires, démonstrations physiques, formation d’agents, distribution de brochures visuelles.  
        - Relation Client & Support : formation du personnel de SAV, mise en place d’une hotline téléphonique, ateliers pratiques, visites régulières sur le terrain.  
        - Partenariats & Négociations : conclure des partenariats avec ONG, coopératives, associations villageoises, négocier des conditions avantageuses avec fournisseurs locaux.  
        Intègre une perspective adaptative : si la demande fluctue, ajuster les stocks, si une nouvelle réglementation apparaît, adapter le produit. Justifie comment chaque activité soutient la proposition de valeur. »

        8. **Partenaires Clés**  
        Méta-Prompt :  
        « Liste et justifie les partenaires stratégiques :  
        - Fournisseurs locaux : garantissant disponibilité et qualité des matières premières.  
        - Distributeurs locaux et intermédiaires informels : accès direct à la clientèle, réduction des coûts d’acquisition.  
        - Partenaires technologiques locaux ou ONG : formation, maintenance, R&D adaptée.  
        - Organismes de certification locaux, influenceurs communautaires, médias (radios, journaux locaux) : augmentent la crédibilité et la confiance.  
        - Institutions financières (microfinance) : faciliter l’accès au crédit, au paiement échelonné.  
        Anticipe les risques (un partenaire clé fait défaut, troubles politiques, pénuries) et prévois des alternatives (autres fournisseurs, diversification géographique). Explique comment ces partenariats renforcent la proposition de valeur et l’efficacité opérationnelle. »

        9. **Structure de Coûts**  
        Méta-Prompt :  
        « Détaille les coûts :  
        - Coûts fixes : salaires des agents locaux, loyers de petits entrepôts, licences minimales, amortissement de matériel de base.  
        - Coûts variables : achat des matières premières, commission aux revendeurs, campagnes radio, formation continue, SAV.  
        - Coûts liés à l’innovation : R&D pour adapter le produit, formation des équipes, tests terrain.  
        Analyse la rentabilité :  
        - Le modèle de revenus couvre-t-il ces coûts ?  
        - Possibilités de réduire les coûts (sourcing local moins cher, économies d’échelle, recyclage, revente de pièces usagées).  
        - Stratégies pour faire face aux fluctuations (augmenter la part de services, moduler les prix, limiter le stock).  
        Explique comment la structure de coûts reste en ligne avec la proposition de valeur, le niveau de vie local, et comment elle assure la pérennité financière à long terme. »

        ---

        **Instructions Finales** :  
        Après avoir utilisé ces méta-prompts pour chaque bloc du Business Model Canvas, effectue une dernière vérification :  
        - Assure-toi que tous les blocs sont cohérents et alignés avec la proposition de valeur, le parcours client et les réalités locales.  
        - Vérifie que l’innovation (low-tech ou high-tech adaptée) est réellement adoptable par la persona, apporte un avantage concurrentiel durable, et que les contraintes (culturelles, économiques, réglementaires, infrastructurelles) sont prises en compte.  
        - Contrôle la rentabilité, la viabilité à long terme, et la flexibilité face aux changements (variations saisonnières, crises économiques, évolution des réglementations ou de la pénétration technologique).  
        - Ajuste les éléments (segments, prix, canaux, partenariats) si nécessaire pour améliorer la robustesse du modèle.  
        - Fournis un récapitulatif global du Business Model, mettant en avant la logique, la cohérence, la proposition de valeur différenciante et quelques chiffres (taille du marché estimée, coûts, revenus, marge, etc.) pour valider la viabilité économique.

        Le résultat final doit être un Business Model clair, complet, adapté au contexte local, prêt à être testé ou implémenté, avec une feuille de route pour l’adoption progressive de l’innovation et une vision claire des points de différenciation face aux solutions traditionnelles ou informelles existantes.
        """,
        
        
        "Startup": """ Tu es un assistant expert en stratégie d’entreprise, marketing, UX, innovation et élaboration de Business Models. Ton rôle est de générer un Business Model complet, clair, chiffré, cohérent et innovant, en suivant trois phases : Configuration Initiale, Étapes Intermédiaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation) et Production Finale (Business Model Canvas).

        Tout au long du processus, tu dois :
        - Prendre en compte la persona (données démographiques, comportementales, capacités d’adoption de l’innovation).
        - Analyser le marché (taille, segments, offres existantes, niveau de satisfaction, tendances).
        - Intégrer les facteurs limitants (technologiques, économiques, culturels, psychologiques, physiologiques, réglementaires).
        - Évaluer la concurrence et comprendre le niveau de satisfaction actuel.
        - Comprendre le parcours client (avant, pendant, après), la carte d’empathie, les gains et souffrances.
        - Vérifier systématiquement la cohérence, proposer des optimisations et ajustements.
        - Avant d’introduire une innovation, t’assurer que la persona est prête à l’adopter.
        - Produire un Business Model Canvas complet (9 blocs), avec des meta-prompts spécifiques pour chacun des blocs.

        Voici les étapes :

        ### Phase 1 : Configuration Initiale (Entrée de Données)

        1. Recueille et structure les informations suivantes :
        - **Persona :**
            - Données démographiques : Âge, sexe, localisation, niveau d’éducation, profession, revenu.
            - Paramètres comportementaux : Sensibilité au prix, budget, fréquence et volume d’achat, perception de la qualité, utilisation technologique, accessibilité, temps disponible, besoins, motivations.
            - Capacité d’adoption de l’innovation : Familiarité technologique, ouverture au changement, barrières psychologiques ou culturelles.
        - **Analyse du Marché :**
            - Taille du marché, segments, valeur totale.
            - Offres concurrentes, niveau de satisfaction, tendances, innovations, comportements émergents.
        - **Facteurs Limitants :**
            - Contraintes technologiques, économiques, culturelles, réglementaires, physiologiques, psychologiques.
        
        Après avoir recueilli ces données, effectue une première analyse critique :
        - Vérifie la cohérence des informations.
        - Identifie les lacunes.
        - Propose des compléments ou ajustements pour optimiser la qualité des données.

        ### Phase 2 : Étapes Intermédiaires (Analyse, Contexte, Empathie, Parcours Client, Optimisation)

        2. **Analyse du Parcours Client & Carte d’Empathie :**
        - Décris le parcours client (avant, pendant, après consommation).
        - Identifie les points de contact, obstacles, moments de vérité, frustrations.
        - Intègre les contraintes physiologiques, psychologiques, économiques, culturelles, technologiques, réglementaires.
        - Crée une carte d’empathie (pensées, sentiments, actions) pour comprendre l’expérience du client à chaque étape.

        3. **Gains et Souffrances :**
        - À partir du parcours client et de la carte d’empathie, liste les gains (bénéfices, réassurance, sentiment de compétence) et les souffrances (problèmes non résolus, frustrations, coûts d’opportunité).

        4. **Élaboration de la Carte de Valeur :**
        - Définis la mission de consommation principale (besoin fondamental).
        - Identifie les gains déjà fournis par les offres actuelles.
        - Mets en évidence les souffrances non adressées.
        - Esquisse une proposition de valeur préliminaire, adaptée à la capacité d’adoption de l’innovation par la persona.

        5. **Détermination du Segment de Clients :**
        - Choisis le type de relation (B2C, B2B, B2B2C…).
        - Priorise les segments (taille, pouvoir d’achat, sensibilité au prix, ouverture à l’innovation, contraintes) qui correspondent le mieux à la proposition de valeur.

        6. **Analyse des Problèmes et Solutions (Canvas de Problème) :**
        - Identifie clairement les problèmes majeurs à résoudre.
        - Associe chaque problème à une solution spécifique, justifie en quoi elle est meilleure que les offres existantes.

        Après ces étapes, effectue une analyse intermédiaire :
        - Vérifie la cohérence du contexte, du parcours client, des solutions proposées.
        - Assure-toi que les innovations sont compréhensibles, utiles et adoptables par la persona.
        - Propose des ajustements stratégiques (simplification de l’offre, ajustement du prix, sélection de segments plus pertinents, etc.) si nécessaire.

        ### Phase 3 : Production Finale du Business Model (Business Model Canvas)

        Sur la base des analyses précédentes, génère un Business Model Canvas complet. Utilise les meta-prompts suivants pour chaque bloc :

        1. **Segments de Clients**  
        Méta-Prompt :  
        « Définis précisément les segments de clients ciblés, en tenant compte :  
        - De leurs caractéristiques sociodémographiques (âge, sexe, localisation, niveau d’éducation, profession, revenu).  
        - De leurs comportements d’achat (fréquence, volume, sensibilité au prix, critères de qualité) et de leur maturité technologique (utilisation d’outils numériques, appareils connectés, plateformes en ligne).  
        - De leur capacité d’adoption de l’innovation (ouverture au changement, barrières psychologiques, éventuelle réticence culturelle).  
        - De leurs contraintes physiologiques (accessibilité, ergonomie), psychologiques (stress, anxiété, besoin de réassurance), économiques (pouvoir d’achat, rapport qualité/prix), culturelles (normes, tabous) et réglementaires (normes légales, certifications).  
        Intègre également des scénarios évolutifs :  
        - Si la technologie évolue, comment ce segment réagit-il ?  
        - S’il y a une crise économique, ces clients réduisent-ils leur consommation ?  
        - Une partie du segment est-elle prête à payer plus pour des options premium ?  
        Justifie pourquoi ces segments sont retenus, comment ils se distinguent de segments non ciblés, et comment leur potentiel de rentabilité, leur facilité d’accès, et leur réceptivité à la proposition de valeur justifient leur inclusion. »

        2. **Proposition de Valeur**  
        Méta-Prompt :  
        « Détaille la proposition de valeur en explicitant :  
        - Les besoins fondamentaux adressés (mission de consommation principale).  
        - Les souffrances clientes (manque de temps, complexité, mauvaise qualité, manque de confiance, crainte face à l’innovation) et comment elles sont résolues.  
        - Les gains fournis (gain de temps, économie d’argent, facilité d’utilisation, statut social, tranquillité d’esprit), y compris les bénéfices émotionnels et symboliques.  
        - La différenciation par rapport aux offres concurrentes (qualité supérieure, innovation plus accessible, prix compétitifs, service client exemplaire, partenariats de prestige).  
        - L’intégration de l’innovation : montre comment elle est introduite progressivement, comment l’éducation ou la formation du client est assurée, et comment les barrières à l’adoption sont levées (essais gratuits, démonstrations, tutoriels, certifications reconnues).  
        - Prévois des variantes de proposition de valeur en fonction des segments, si nécessaire (une version premium pour les early adopters innovants, une version simplifiée pour les plus conservateurs). »

        3. **Canaux de Distribution**  
        Méta-Prompt :  
        « Définis les canaux par lesquels les clients seront informés, convaincus, achèteront et utiliseront le produit/service. Considère :  
        - Les canaux en ligne (site web, application mobile, plateformes e-learning, réseaux sociaux, partenariats avec marketplaces, influenceurs, SEO, SEA).  
        - Les canaux hors ligne (magasins physiques, salons professionnels, conférences, revendeurs, agents sur le terrain).  
        - La nécessité de cohérence entre les points de contact (omnicanal), la simplicité d’accès, le besoin d’accompagnement pédagogique (webinaires, tutoriels vidéo), et les contraintes technologiques de la persona (faible bande passante, préférence pour un canal mobile vs desktop).  
        - L’adaptabilité des canaux si les conditions du marché changent (pénurie d’un canal, évolution légale, concurrence d’un nouveau canal).  
        Justifie pourquoi chaque canal est choisi, comment il s’intègre dans le parcours client, comment il favorise l’adoption de l’innovation, et comment il est optimisé pour réduire les coûts d’acquisition et améliorer la satisfaction. »

        4. **Relations Clients**  
        Méta-Prompt :  
        « Décris la nature et la qualité des relations que l’entreprise établira avec ses clients :  
        - Personnalisation : existe-t-il un accompagnement individuel, des conseils sur mesure, une assistance humaine ou une IA conversationnelle ?  
        - Communauté : les clients peuvent-ils interagir entre eux (forums, réseaux sociaux, clubs, rencontres physiques) pour renforcer leur sentiment d’appartenance et échanger des expériences ?  
        - Automatisation : y a-t-il des éléments de self-service, de chatbots, de bases de connaissances en ligne ? Est-ce adapté aux cibles moins technophiles ?  
        - Fidélisation : cartes de fidélité, programmes de récompenses, contenus exclusifs, mises à jour gratuites, offres spéciales pour clients fidèles.  
        - Gestion des plaintes et retours : procédures de remboursement, garantie de satisfaction, SLA pour répondre aux demandes critiques.  
        Intègre la dimension psychologique (rassurer les clients sur l’innovation), culturelle (certains clients préfèrent un contact humain), réglementaire (besoin de conformité avec les lois sur la protection des données).  
        Explique comment ces relations évoluent au fil du temps (du premier contact à la fidélisation), comment elles améliorent la CLV, et comment elles s’adaptent aux changements de marché (nouveaux concurrents, crises économiques). »

        5. **Sources de Revenus**  
        Méta-Prompt :  
        « Détaille les mécanismes de génération de revenus :  
        - Modèle de tarification : abonnement mensuel, paiement à l’usage, achat unique, freemium avec options premium, licences, commissions.  
        - Justification des prix : comment le prix reflète-t-il la valeur perçue par le client ? Est-il aligné avec le pouvoir d’achat du segment, la concurrence, la qualité et l’innovation proposée ?  
        - Options de réduction des freins économiques : essais gratuits, garantie satisfait ou remboursé, paiement échelonné, remises pour les early adopters.  
        - Diversification des revenus : ventes croisées, upselling, partenariats, publicité, formation complémentaire, monétisation de données (en respectant la réglementation).  
        - Adaptation à des changements de contexte : si le marché se contracte, proposer un modèle plus flexible ? Si une réglementation limite certains types de revenus, anticiper une alternative ?  
        Explique comment cette structure de revenus soutient la croissance, la rentabilité, et s’intègre avec les coûts prévus. Vérifie la cohérence avec la proposition de valeur et la sensibilité au prix de la persona. »

        6. **Ressources Clés**  
        Méta-Prompt :  
        « Identifie toutes les ressources indispensables :  
        - Ressources Humaines : équipes multidisciplinaires (ingénieurs, designers UX, experts marketing, formateurs, support client multilingue) nécessaires à la création, maintenance, amélioration de l’offre.  
        - Ressources Technologiques : plateformes e-learning, serveurs, logiciels de personnalisation, outils d’IA, applications mobiles, infrastructure IT sécurisée.  
        - Ressources Intellectuelles : brevets, marques, contenus propriétaires, méthodologies exclusives, licences de tiers, données clients protégées.  
        - Ressources Financières : capitaux nécessaires au lancement, trésorerie pour résister à une période de faible demande, fonds pour R&D.  
        - Ressources Relationnelles : partenariats stratégiques, accès à un réseau d’influenceurs, certification par des organismes reconnus.  
        Explique pour chaque ressource pourquoi elle est critique, comment elle se combine avec les autres pour délivrer la proposition de valeur, soutenir l’adoption de l’innovation, et maintenir un avantage concurrentiel. Prends en compte la robustesse de la chaîne d’approvisionnement, la résilience face aux crises, et la propriété intellectuelle. »

        7. **Activités Clés**  
        Méta-Prompt :  
        « Décris les activités indispensables pour que le Business Model fonctionne :  
        - Développement & Innovation : R&D, amélioration continue, intégration de nouvelles fonctionnalités, veille concurrentielle, tests utilisateurs.  
        - Production & Livraison : création de contenu, mise à jour régulière, gestion du stock (si produit physique), maintenance technique, logistique.  
        - Marketing & Ventes : campagnes publicitaires, référencement, webinaires de démonstration, éducation du marché, gestion des promotions.  
        - Relation Client & Support : formation du personnel du support, chatbots, assistance multicanal, traitement des plaintes, suivi de la satisfaction.  
        - Partenariats & Négociations : recherche, signature et entretien des partenariats clés, mise en place de conditions avantageuses.  
        Intègre une perspective adaptative :  
        - Quelles activités mener si la demande fluctue fortement ?  
        - Comment réallouer les ressources si une nouvelle réglementation émerge ?  
        Justifie comment chaque activité soutient la proposition de valeur, favorise l’adoption de l’innovation, et contribue à la rentabilité globale. »

        8. **Partenaires Clés**  
        Méta-Prompt :  
        « Liste et justifie les partenaires stratégiques critiques :  
        - Fournisseurs : apportant des ressources rares, de haute qualité ou à un coût avantageux.  
        - Distributeurs : offrant un accès facilité à certains segments, réduisant les coûts d’acquisition, améliorant la visibilité.  
        - Partenaires technologiques : fournissant une infrastructure fiable, des outils d’IA performants, ou des solutions complémentaires (API, intégrations).  
        - Organismes de certification, influenceurs, médias spécialisés : augmentant la crédibilité, validant la qualité, rassurant sur l’innovation.  
        - Associations professionnelles, clusters, écosystèmes sectoriels : permettant de suivre les tendances, d’anticiper les changements réglementaires, d’échanger les bonnes pratiques.  
        Explique comment ces partenariats renforcent la proposition de valeur, améliorent la confiance du client, augmentent l’efficacité opérationnelle, réduisent les coûts ou les risques, et soutiennent la stratégie à long terme. Anticipe les risques : et si un partenaire clé fait défaut ? Quels sont les plans B ? »

        9. **Structure de Coûts**  
        Méta-Prompt :  
        « Détaille tous les coûts engendrés par les ressources, activités et partenariats clés :  
        - Coûts fixes (salaires, loyers, licences, amortissement de l’infrastructure).  
        - Coûts variables (marketing, support client, acquisition de nouveaux outils, commission aux partenaires).  
        - Coûts liés à l’innovation (R&D, tests, formations du personnel), et comment ils sont amortis dans le temps.  
        Analyse la rentabilité :  
        - Le modèle de revenus couvre-t-il ces coûts ?  
        - Quelles mesures de réduction de coûts sont possibles (automatisation, sourcing moins cher, économies d’échelle) ?  
        - Comment réagir face à des fluctuations du marché (baisse de la demande, hausse des prix des ressources) ?  
        Explique comment la structure de coûts reste alignée avec la proposition de valeur, les segments, et les moyens financiers de l’entreprise. Justifie la pérennité financière en montrant que les marges sont satisfaisantes, que le CAC est raisonnable par rapport à la CLV, et que le modèle reste rentable même en cas de stress. »

        ### Instructions Finales

        Après avoir utilisé ces méta-prompts pour chaque bloc du Business Model Canvas, effectue une dernière vérification :

        - Assure-toi que tous les blocs sont cohérents entre eux et s’alignent parfaitement avec la proposition de valeur et le parcours client.
        - Vérifie que l’innovation proposée est bien adoptable par la persona, qu’elle apporte un avantage concurrentiel durable, et que les contraintes sont gérées.  
        - Contrôle la rentabilité, la viabilité à long terme, et la flexibilité pour s’adapter aux changements de marché.
        - Ajuste les éléments (segments, prix, canaux, partenariats) si nécessaire pour améliorer la robustesse du modèle.

        Le résultat final doit être un Business Model clair, complet, et prêt à être testé ou implémenté, avec une feuille de route pour l’adoption de l’innovation et une vision claire des points de différenciation face à la concurrence.


        Enfin, fournis un récapitulatif global du Business Model, mettant en avant la logique, la cohérence, et la proposition de valeur différenciante. Indique, si possible, des chiffres (taille du marché, CAC, CLV, taux de conversion, CA projeté) pour valider la viabilité économique.""",
       
        "Autre": "Fournissez une approche générale adaptée à votre entreprise."
    }
    return metaprompts.get(type_entreprise, metaprompts["Autre"])




def obtenir_business_model(nom_entreprise, type_entreprise, rubriques):
    """
    Interroge ChatGPT (API OpenAI) pour générer le contenu textuel
    des différents blocs du Business Model Canvas.
    'type_entreprise' peut être "PME", "Startup", "Grande Entreprise", etc.
    """
    
    
    # Récupérer le metaprompt basé sur le type d'entreprise
    metaprompt = get_metaprompt(type_entreprise)
    print(rubriques)
    
   
    # Prompt ajusté sans numérotation dans les titres
    prompt = f"""
    {metaprompt}
    
    Mener la reflexions du generation du business modele sur base des indications(Méta-Prompt) precedents du metaprompts; 
    Chercher les chiffres et autres données sur internet, assurer-vous d'etre trop precis et excat en fonction fonction des données collecter sur internet 
    Génère le contenu d'un Business Model Canvas en format HTML pour une entreprise nommée '{nom_entreprise}'.
    Le type d'entreprise est : {type_entreprise}.
    et dont les données complementaires (non obligatoire pour l'utilisateur) pour chaque bloc se trouve dans : {rubriques}.
    si l'utlisateur a donner les données complementaires, veuillez en tenir compte dans la generation, et ca doit etre imperativement prioritaire.
    Si dans un bloque un utilisateur n'as pas donner des informations (elements), veuillez generer,
    Si l'utilisateur à donné des elements que vous juger peu, generer d'autres et les ajoutées à ce que l'utlisateur à fournit.
    
    à faire imperativement est:
    Je veux impérativement 9 blocs distincts, rédigés en français, avec les titres en gras et des listes à puces si nécessaire :
      - Partenaires clés
      - Activités clés
      - Offre (proposition de valeur)
      - Relation client
      - Segments de clientèle
      - Ressources clés
      - Canaux de distribution
      - Structure des coûts
      - Sources de revenus
    Fournis 5 à 10 points ou élements(phrases) , meme plus pour chacun afin d'avoir un contenu riche et adapté, soyez concis.
    """

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Tu es un assistant expert en génération de business plan."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=10000,
            temperature=0.8
        )
        html_genere = response.choices[0].message.content.strip()
        return html_genere
    except Exception as e:
        st.error(f"Erreur lors de la génération du contenu : {e}")
        return ""

# ----------------------------------------------------------------------------
# 2) Fonction pour créer le fichier Word (format tableau) avec python-docx
# ----------------------------------------------------------------------------

def generer_docx_business_model(nom_entreprise, date_bmc, contenu_business_model):
    """
    Construit un document Word reproduisant un tableau avec la disposition souhaitée
    pour le Business Model Canvas. La mise en forme inclut des titres en gras et
    des listes à puces.
    'contenu_business_model' : le contenu HTML renvoyé par ChatGPT,
    qu'on découpe ensuite pour remplir chaque bloc.
    """
    # Créer un nouveau document Word
    doc = Document()

    # Définir les styles de base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Titre principal
    titre = doc.add_heading(level=1)
    titre_run = titre.add_run(f"Business Model Canvas de {nom_entreprise}")
    titre_run.bold = True
    titre.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Date
    date_paragraph = doc.add_paragraph()
    date_run = date_paragraph.add_run(f"Date : {date_bmc}")
    date_run.bold = True
    date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Ajouter un saut de ligne
    doc.add_paragraph("")

    # Créer un tableau de 6 lignes × 5 colonnes
    table = doc.add_table(rows=6, cols=5)
    table.style = 'Table Grid'

    # Ajuster les largeurs des colonnes (en pouces)
    for col in table.columns:
        for cell in col.cells:
            cell.width = Inches(1.8)  # Ajustez selon vos besoins

    # 1) Ligne 0 : Titre (fusion des 5 colonnes)
    cell00 = table.cell(0, 0)
    cell00_merge = cell00.merge(table.cell(0, 4))
    cell00_merge.text = f"Business Model Canvas de {nom_entreprise}"
    for paragraph in cell00_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 2) Ligne 1 : Nom de l'entreprise et Date (fusion des colonnes)
    cell10 = table.cell(1, 0)
    cell10_merge = cell10.merge(table.cell(1, 2))
    cell10_merge.text = f"**Nom de l'entreprise**: {nom_entreprise}"
    for paragraph in cell10_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    cell13 = table.cell(1, 3)
    cell13_merge = cell13.merge(table.cell(1, 4))
    cell13_merge.text = f"**Date**: {date_bmc}"
    for paragraph in cell13_merge.paragraphs:
        for run in paragraph.runs:
            run.bold = True
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # 3) Ligne 2 : Headers des 5 blocs
    headers = ["Partenaires clés", "Activités clés", "Offre (proposition de valeur)", 
               "Relation client", "Segments de clientèle"]
    for idx, header in enumerate(headers):
        cell = table.cell(2, idx)
        paragraphe = cell.paragraphs[0]
        run = paragraphe.add_run(header)
        run.bold = True
        paragraphe.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 4) Ligne 3 : Contenus des 5 blocs
    # Initialiser les blocs
    blocs = {
        "Partenaires clés": "",
        "Activités clés": "",
        "Offre (proposition de valeur)": "",
        "Relation client": "",
        "Segments de clientèle": ""
    }

    # Utiliser BeautifulSoup pour parser le HTML
    soup = BeautifulSoup(contenu_business_model, 'html.parser')

    # Fonction pour trouver le bon header tag (h3 par défaut, avec flexibilité)
    def trouver_header(soup, header):
        # Regex pour capturer optionnellement des numéros suivis de points et espaces
        pattern = rf"^(?:\d+\.\s*)?{re.escape(header)}$"
        # Chercher dans les balises h3
        header_tag = soup.find(['h2', 'h3', 'h4', 'h5', 'h6'], text=re.compile(pattern, re.IGNORECASE))
        return header_tag

    # Extraire chaque bloc
    for header in blocs.keys():
        h_tag = trouver_header(soup, header)
        if h_tag:
            content = []
            for sibling in h_tag.find_next_siblings():
                if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                    break  # Arrêter si un nouveau header est trouvé
                if sibling.name == 'ul':
                    for li in sibling.find_all('li'):
                        content.append(f"- {li.get_text(strip=True)}")
                elif sibling.name == 'p':
                    content.append(sibling.get_text(strip=True))
                elif isinstance(sibling, str):
                    text = sibling.strip()
                    if text:
                        content.append(text)
            blocs[header] = '\n'.join(content)

    # Debug: Afficher les blocs extraits (à désactiver en production)
    # st.write("Blocs extraits :", blocs)

    # Fonction pour ajouter du contenu formaté dans une cellule
    def ajouter_contenu(cell, titre, contenu):
        """
        Ajoute du contenu formaté dans une cellule Word.
        Le titre est en gras, suivi de listes à puces si nécessaire.
        """
        # Supprimer le texte initial (par défaut) dans la cellule
        cell.text = ""

        # Ajouter le titre en gras
        paragraphe = cell.add_paragraph()
        run = paragraphe.add_run(titre)
        run.bold = True

        # Ajouter le contenu
        # Diviser le contenu par les sauts de ligne
        lignes = contenu.split('\n')
        for ligne in lignes:
            ligne = ligne.strip()
            if not ligne:
                continue
            # Vérifier si la ligne commence par '-', '+', '•' pour une liste à puces
            if re.match(r'^[-+•]\s+', ligne):
                # Ajouter une puce
                item = re.sub(r'^[-+•]\s+', '', ligne)
                p = cell.add_paragraph(item, style='List Bullet')
            else:
                # Ajouter un paragraphe normal
                p = cell.add_paragraph(ligne)

    # Remplir les cellules de la ligne 3
    ordre_blocs = [
        "Partenaires clés", "Activités clés", "Offre (proposition de valeur)",
        "Relation client", "Segments de clientèle"
    ]

    for idx, bloc in enumerate(ordre_blocs):
        cell = table.cell(3, idx)
        ajouter_contenu(cell, bloc, blocs[bloc])

    # 5) Ligne 4 : Structure de coûts (fusion 3 cols) et Sources de revenus (fusion 2 cols)
    # Fusionner les cellules pour "Structure de coûts" (colonnes 0-2)
    cell40 = table.cell(4, 0)
    cell40_merge = cell40.merge(table.cell(4, 2))
    cell40_merge.text = f"**Structure de coûts**:\n\n"

    # Fusionner les cellules pour "Sources de revenus" (colonnes 3-4)
    cell43 = table.cell(4, 3)
    cell43_merge = cell43.merge(table.cell(4, 4))
    cell43_merge.text = f"**Sources de revenus**:\n\n"

    # Extraire les contenus pour ces blocs
    structure_couts = ""
    sources_revenus = ""

    # Structure des coûts
    strong_tag = trouver_header(soup, "Structure des coûts")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        structure_couts = '\n'.join(content)

    # Sources de revenus
    strong_tag = trouver_header(soup, "Sources de revenus")
    if strong_tag:
        content = []
        for sibling in strong_tag.find_next_siblings():
            if sibling.name and re.match(r'^h[2-6]$', sibling.name, re.IGNORECASE):
                break
            if sibling.name == 'ul':
                for li in sibling.find_all('li'):
                    content.append(f"- {li.get_text(strip=True)}")
            elif sibling.name == 'p':
                content.append(sibling.get_text(strip=True))
            elif isinstance(sibling, str):
                text = sibling.strip()
                if text:
                    content.append(text)
        sources_revenus = '\n'.join(content)

    # Remplir les cellules fusionnées
    ajouter_contenu(cell40_merge, "Structure de coûts", structure_couts)
    ajouter_contenu(cell43_merge, "Sources de revenus", sources_revenus)

    # Ajuster les paragraphes existants
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

    # Ajouter un saut de ligne à la fin
    doc.add_paragraph("")

    # Convertir le document en binaire pour téléchargement via Streamlit
    fichier_io = BytesIO()
    doc.save(fichier_io)
    fichier_io.seek(0)
    return fichier_io

# ----------------------------------------------------------------------------
# 3) Application Streamlit
# ----------------------------------------------------------------------------
def main():
    st.title("Générateur de Business Model Canvas")
    st.write(
        "Cette application génère automatiquement un Business Model Canvas (format Word) "
        "en se basant sur le nom et le type de votre entreprise, via ChatGPT."
    )
    
        # Champs de saisie
    nom_entreprise = st.text_input("Nom de l'entreprise", value="COIFFURE MOBILE S.a.r.l")
    type_entreprise = st.selectbox("Type d'entreprise", ["PME", "Startup", "Autre"])
    date_bmc = st.date_input("Date du BMC", value=datetime.date.today())
    
     # Formulaire pour les 9 rubriques du BMC
        # Variable pour contrôler l'état global des expanders
    expand_all = st.checkbox("Étendre / Réduire tout le formulaire", value=False)

    # Formulaire pour les 9 rubriques du BMC
    st.header("Informations Complémentaires pour le Business Model Canvas")

    with st.form("form_rubriques"):
        st.write("Veuillez remplir les informations pour chaque rubrique du Business Model Canvas.")

        with st.expander("Partenaires clés", expanded=expand_all):
            st.markdown("""
            **Partenaires clés :**  
            Identifiez les organisations ou individus essentiels à votre activité. Par exemple :  
            - Fournisseurs  
            - Banques ou institutions financières  
            - Partenaires stratégiques  
            - Associations ou gouvernements locaux  
            """)
            partenaire_cles = st.text_area("Listez vos principaux partenaires.", key="partenaires_cles")

        with st.expander("Activités clés", expanded=expand_all):
            st.markdown("""
            **Activités clés :**  
            Décrivez les tâches ou processus les plus importants pour exécuter votre modèle économique. Par exemple :  
            - Fabrication de produits  
            - Marketing et vente  
            - Recherche et développement  
            - Gestion des relations avec les clients  
            """)
            activites_cles = st.text_area("Décrivez vos activités principales.", key="activites_cles")

        with st.expander("Offre (proposition de valeur)", expanded=expand_all):
            st.markdown("""
            **Proposition de valeur :**  
            Expliquez ce que vous offrez à vos clients et ce qui vous différencie de vos concurrents. Par exemple :  
            - Résolution d'un problème spécifique  
            - Amélioration d'un besoin existant  
            - Caractéristiques uniques de vos produits ou services  
            """)
            offre_valeur = st.text_area("Décrivez votre proposition de valeur.", key="offre_valeur")

        with st.expander("Relation client", expanded=expand_all):
            st.markdown("""
            **Relation client :**  
            Décrivez comment vous interagissez avec vos clients. Par exemple :  
            - Assistance personnalisée  
            - Automatisation des services (chatbots, self-service)  
            - Programmes de fidélisation  
            """)
            relation_client = st.text_area("Décrivez comment vous gérez vos relations clients.", key="relation_client")

        with st.expander("Segments de clientèle", expanded=expand_all):
            st.markdown("""
            **Segments de clientèle :**  
            Identifiez vos différents groupes de clients cibles. Par exemple :  
            - Particuliers (par revenus, âge, localisation)  
            - Entreprises (par secteur ou taille)  
            - Marchés de niche  
            """)
            segments_clientele = st.text_area("Définissez vos segments de clientèle.", key="segments_clientele")

        with st.expander("Ressources clés", expanded=expand_all):
            st.markdown("""
            **Ressources clés :**  
            Listez les ressources nécessaires pour exécuter vos activités. Par exemple :  
            - Ressources physiques (locaux, machines)  
            - Ressources humaines (compétences clés, équipes)  
            - Ressources financières (fonds, prêts)  
            """)
            ressources_cles = st.text_area("Listez vos ressources principales.", key="ressources_cles")

        with st.expander("Canaux de distribution", expanded=expand_all):
            st.markdown("""
            **Canaux de distribution :**  
            Décrivez comment vos produits ou services atteignent vos clients. Par exemple :  
            - Boutiques physiques  
            - Plateformes en ligne  
            - Distributeurs tiers  
            """)
            canaux_distribution = st.text_area("Décrivez vos canaux de distribution.", key="canaux_distribution")

        with st.expander("Structure de coûts", expanded=expand_all):
            st.markdown("""
            **Structure de coûts :**  
            Énumérez les principaux coûts liés à votre activité. Par exemple :  
            - Coûts de production  
            - Salaires et charges sociales  
            - Dépenses marketing et publicitaires  
            """)
        
            structure_couts = st.text_area("Décrivez votre structure de coûts.", key="structure_couts")

        with st.expander("Sources de revenus", expanded=expand_all):
            st.markdown("""
            **Sources de revenus :**  
            Décrivez comment vous générez des revenus. Par exemple :  
            - Vente de produits ou services  
            - Abonnements  
            - Publicité ou partenariats  
            """)
            sources_revenus = st.text_area("Décrivez vos sources de revenus.", key="sources_revenus")

        submit_form = st.form_submit_button("Valider les Informations")
             # Récupérer les informations du formulaire
        if submit_form:
                # Récupérer les informations du formulaire
            rubriques = {
                "Partenaires clés": partenaire_cles,
                "Activités clés": activites_cles,
                "Offre (proposition de valeur)": offre_valeur,
                "Relation client": relation_client,
                "Segments de clientèle": segments_clientele,
                "Ressources clés": ressources_cles,
                "Canaux de distribution": canaux_distribution,
                "Structure de coûts": structure_couts,
                "Sources de revenus": sources_revenus
                }
        else:
                # Si le formulaire n'est pas soumis, utiliser des valeurs vides
            rubriques = {
                "Partenaires clés": "",
                "Activités clés": "",
                "Offre (proposition de valeur)": "",
                "Relation client": "",
                "Segments de clientèle": "",
                "Ressources clés": "",
                "Canaux de distribution": "",
                "Structure de coûts": "",
                "Sources de revenus": ""
                }    



    # Bouton pour générer
    if st.button("Générer le Business Model Canvas"):
        with st.spinner("Génération en cours..."):
            
            rubriques = {
                "Partenaires clés": partenaire_cles,
                "Activités clés": activites_cles,
                "Offre (proposition de valeur)": offre_valeur,
                "Relation client": relation_client,
                "Segments de clientèle": segments_clientele,
                "Ressources clés": ressources_cles,
                "Canaux de distribution": canaux_distribution,
                "Structure de coûts": structure_couts,
                "Sources de revenus": sources_revenus
             }
            print(rubriques)
            # 1) Obtenir le contenu textuel via ChatGPT
            contenu_bmc = obtenir_business_model(nom_entreprise, type_entreprise, rubriques)
            

            if not contenu_bmc:
                st.error("Erreur lors de la génération du contenu. Veuillez réessayer.")
                return

            # 2) Générer le document Word en mémoire
            docx_bytes = generer_docx_business_model(
                nom_entreprise=nom_entreprise,
                date_bmc=date_bmc.strftime("%d %B %Y"),
                contenu_business_model=contenu_bmc
            )

        st.success("Business Model Canvas généré avec succès !")

        # Proposer le téléchargement du document Word
        st.download_button(
            label="Télécharger le Business Model Canvas (Word)",
            data=docx_bytes,
            file_name=f"BMC_{nom_entreprise.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Optionnel : Afficher le contenu généré pour vérification
        st.subheader("Contenu Généré par ChatGPT")
        st.markdown(contenu_bmc, unsafe_allow_html=True)

# Point d'entrée
if __name__ == "__main__":
    main()